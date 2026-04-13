import os
import re
import json
from io import BytesIO
from typing import List, Tuple
from datetime import date

import streamlit as st

try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    docx = None
    Document = None
    Pt = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False


st.set_page_config(page_title="Fresher Job Cracker AI", page_icon="💼", layout="wide")

st.markdown(
    """
    <style>
    :root {
        --bg: #0f1115;
        --card: #171a21;
        --card2: #1f2430;
        --text: #f8fafc;
        --muted: #b8c0cc;
        --line: rgba(255,255,255,0.08);
        --accent: #fc8019;
        --accent2: #ff9a4f;
        --danger: #ef4444;
        --success: #22c55e;
        --shadow: 0 14px 34px rgba(0,0,0,0.24);
    }
    .stApp {
        background: radial-gradient(circle at top right, rgba(252,128,25,0.10), transparent 22%), linear-gradient(180deg, #0c0f14 0%, #101318 100%);
        color: var(--text);
    }
    .block-container { max-width: 1250px; padding-top: 1.2rem; padding-bottom: 2rem; }
    [data-testid="stSidebar"] { background: linear-gradient(180deg,#141821,#10141a); border-right: 1px solid var(--line); }
    .hero {
        background: linear-gradient(135deg, rgba(252,128,25,0.24), rgba(252,128,25,0.08));
        border: 1px solid rgba(252,128,25,0.18);
        border-radius: 26px; padding: 1.35rem 1.45rem; box-shadow: var(--shadow); margin-bottom: 1rem;
    }
    .hero-title { font-size: 2.3rem; font-weight: 800; letter-spacing: -0.03em; color: white; margin-bottom: 0.35rem; }
    .hero-sub { font-size: 1rem; color: #fff1e6; line-height: 1.6; max-width: 950px; }
    .chip-row { display:flex; flex-wrap:wrap; gap:0.55rem; margin-top:0.9rem; }
    .chip { padding:0.45rem 0.8rem; border-radius:999px; background:rgba(255,255,255,0.08); color:white; font-size:0.86rem; font-weight:700; border:1px solid rgba(255,255,255,0.08); }
    .soft-card { background: linear-gradient(180deg, rgba(255,255,255,0.03), rgba(255,255,255,0.015)); border: 1px solid var(--line); border-radius: 22px; padding: 1rem; box-shadow: var(--shadow); margin-bottom: 1rem; }
    .section-title { font-size: 1.45rem; font-weight: 800; color: white; margin-bottom: 0.3rem; }
    .section-sub { color: var(--muted); font-size: 0.95rem; margin-bottom: 0.4rem; }
    .step-grid { display:grid; grid-template-columns: repeat(4, minmax(0,1fr)); gap:0.75rem; margin: 0.8rem 0 1rem 0; }
    .step { background: linear-gradient(180deg,#171c24,#131821); border:1px solid var(--line); border-radius:18px; padding:0.95rem; min-height:112px; }
    .step-num { color: var(--accent2); text-transform: uppercase; letter-spacing: 0.06em; font-size:0.78rem; font-weight: 800; margin-bottom: 0.35rem; }
    .step-head { color:white; font-weight:800; font-size:1rem; margin-bottom:0.2rem; }
    .step-copy { color:var(--muted); font-size:0.88rem; line-height:1.45; }
    .paywall { background: linear-gradient(135deg, rgba(252,128,25,0.16), rgba(255,255,255,0.03)); border: 1px solid rgba(252,128,25,0.28); border-radius: 22px; padding: 1.1rem; box-shadow: var(--shadow); }
    .pay-title { font-size:1.25rem; font-weight:800; color:white; margin-bottom:0.2rem; }
    .pay-price { font-size:2rem; font-weight:800; color:white; margin:0.35rem 0; }
    .blur-box { position:relative; border:1px dashed rgba(255,255,255,0.14); border-radius:18px; padding:1rem; margin-top:0.7rem; overflow:hidden; }
    .blur-inner { filter: blur(4px); opacity:0.55; user-select:none; }
    .lock-overlay { position:absolute; inset:0; display:flex; align-items:center; justify-content:center; background: linear-gradient(180deg, rgba(10,12,16,0.08), rgba(10,12,16,0.34)); color:white; font-weight:800; font-size:1rem; }
    .warning-card { border:1px solid rgba(239,68,68,0.25); background: rgba(239,68,68,0.07); border-radius:18px; padding:0.9rem; }
    .success-card { border:1px solid rgba(34,197,94,0.25); background: rgba(34,197,94,0.07); border-radius:18px; padding:0.9rem; }
    div[data-testid="stMetric"] { background: linear-gradient(180deg,#181d25,#131720); border:1px solid var(--line); border-radius:18px; padding:0.75rem 0.9rem; }
    div[data-testid="stMetricValue"] { font-size:2rem; font-weight:800; color:white; }
    .stButton > button, .stDownloadButton > button {
        background: linear-gradient(135deg, var(--accent), #ff9a4f) !important;
        color: white !important; border:none !important; border-radius:14px !important; min-height:2.9rem !important; font-weight:800 !important;
        box-shadow: 0 8px 22px rgba(252,128,25,0.24);
    }
    .stButton > button:hover, .stDownloadButton > button:hover { transform: translateY(-1px); }
    [data-baseweb="tab"] { background:#171c24 !important; border:1px solid var(--line) !important; border-radius:999px !important; padding:0.52rem 0.92rem !important; }
    [aria-selected="true"][data-baseweb="tab"] { background: linear-gradient(135deg,var(--accent),#ff9a4f) !important; border-color:transparent !important; }
    .small { color:var(--muted); font-size:0.88rem; line-height:1.5; }
    @media (max-width: 1080px) { .step-grid { grid-template-columns: repeat(2, minmax(0,1fr)); } }
    @media (max-width: 640px) { .step-grid { grid-template-columns:1fr; } .hero-title { font-size:1.75rem; } }
    </style>
    """,
    unsafe_allow_html=True,
)


# ---------- Helpers ----------
def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def extract_text_from_upload(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    file_name = (uploaded_file.name or "").lower()
    try:
        if file_name.endswith(".txt"):
            return uploaded_file.read().decode("utf-8", errors="ignore")
        if file_name.endswith(".docx") and docx is not None:
            document = docx.Document(uploaded_file)
            return "\n".join([p.text for p in document.paragraphs if p.text.strip()])
        if file_name.endswith(".pdf") and PdfReader is not None:
            reader = PdfReader(uploaded_file)
            return "\n".join([(page.extract_text() or "") for page in reader.pages])
    except Exception:
        return ""
    return ""


def extract_keywords(text: str) -> List[str]:
    words = re.findall(r"[a-zA-Z][a-zA-Z\-\+\.]{1,}", clean_text(text).lower())
    stopwords = {"the","and","for","with","you","your","our","are","will","who","that","this","from","have","has","had","job","role","work","team","into","all","any","but","not","can","able","using","use","their","they","them","what","when","where","how","why","about","should","must","good","strong","such","looking","candidate","candidates","skills","skill","experience","preferred","required","years","year","month","months","day","days","new","best","need","plus","get","build","make","help","across","within","ability","including","responsible","knowledge","understanding","seeking","fresher","graduate"}
    filtered = [w for w in words if len(w) > 2 and w not in stopwords]
    freq = {}
    for w in filtered:
        freq[w] = freq.get(w, 0) + 1
    return [w for w, _ in sorted(freq.items(), key=lambda x: (-x[1], x[0]))[:30]]


def estimate_match(resume_text: str, jd_text: str) -> Tuple[int, List[str], List[str]]:
    resume_keywords = set(extract_keywords(resume_text))
    jd_keywords = set(extract_keywords(jd_text))
    if not jd_keywords:
        return 0, [], []
    overlap = sorted(list(resume_keywords & jd_keywords))
    missing = sorted(list(jd_keywords - resume_keywords))[:12]
    score = int((len(overlap) / max(len(jd_keywords), 1)) * 100)
    score = max(15, min(score, 95)) if resume_text.strip() and jd_text.strip() else 0
    return score, overlap[:12], missing


def to_bullets(block: str, default_line: str) -> str:
    lines = [line.strip("-• ") for line in (block or "").split("\n") if line.strip()]
    if not lines:
        lines = [default_line]
    return "\n".join([f"- {line}" for line in lines])


def build_resume_fallback(name: str, email: str, phone: str, location: str, target_role: str, profile_type: str,
                          summary: str, education: str, experience: str, skills: str, projects: str, certifications: str) -> str:
    defaults = {
        "No experience": f"Motivated fresher targeting {target_role} roles with strong learning agility, communication skills, and structured problem-solving.",
        "Internship only": f"Early-career candidate targeting {target_role} roles with internship exposure, business understanding, and the ability to execute in fast-paced environments.",
        "Career switch": f"Professional transitioning into {target_role} roles with transferable skills, adaptability, and a strong willingness to learn and deliver results.",
        "Experienced": f"Results-oriented candidate targeting {target_role} roles with relevant execution experience, stakeholder handling, and problem-solving skills.",
    }
    return f"""# {name}
{email} | {phone} | {location}

## Target Role
{target_role}

## Profile Type
{profile_type}

## Professional Summary
{summary or defaults.get(profile_type, defaults['No experience'])}

## Education
{education or 'Add your latest degree, institute, and graduation year'}

## Experience
{to_bullets(experience, 'Add internship, campus leadership, freelance work, or responsibility-based experience relevant to the target role')}

## Projects
{to_bullets(projects, 'Add one role-relevant project showing analysis, execution, research, sales, or operations impact')}

## Certifications
{certifications or 'Add relevant certifications here'}

## Skills
{skills or 'Communication, Excel, PowerPoint, Research, Problem Solving, Stakeholder Management'}
"""


def build_cover_letter_fallback(name: str, company: str, role: str, summary: str, strengths: str, jd_text: str) -> str:
    keywords = extract_keywords(jd_text)[:6]
    kw_text = ", ".join(keywords) if keywords else "communication, execution, and adaptability"
    return f"""Dear Hiring Manager,

I am writing to express my interest in the {role} role at {company}. I am excited about this opportunity because it aligns with my background and my interest in building a strong career in this field.

{summary or 'I bring a combination of business understanding, learning agility, and a strong willingness to take ownership.'} My strengths include {strengths or 'communication, structured thinking, and execution'}. Based on the job description, I understand that the role values {kw_text}. I believe these expectations align well with my profile and potential.

I would welcome the chance to discuss how I can contribute to your team. Thank you for your time and consideration.

Sincerely,
{name}
"""


def score_resume(resume_text: str, jd_text: str, profile_type: str, skills_text: str, experience_text: str,
                 projects_text: str, target_role: str) -> Tuple[int, List[str], List[str], List[str]]:
    score = 35
    feedback, positives, suggestions = [], [], []
    skills_count = len([s for s in skills_text.split(",") if s.strip()])
    exp_lines = len([x for x in experience_text.split("\n") if x.strip()])
    proj_lines = len([x for x in projects_text.split("\n") if x.strip()])
    has_numbers = bool(re.search(r"\d", resume_text or ""))
    has_summary = "## Professional Summary" in resume_text and len(clean_text(resume_text.split("## Professional Summary")[-1][:220])) > 20
    if has_summary:
        score += 8
        positives.append("You have a usable summary section.")
    else:
        feedback.append("Your resume lacks a clear professional summary.")
    if skills_count >= 5:
        score += 10
        positives.append("Your skills section has decent coverage.")
    else:
        feedback.append("Your skills section is too thin and may look weak to recruiters.")
    if exp_lines >= 2:
        score += 12
        positives.append("You have enough experience or internship bullets to work with.")
    else:
        feedback.append("You need stronger experience bullets, even if they come from internships, college work, events, or responsibilities.")
    if proj_lines >= 1:
        score += 10
        positives.append("You have at least one project, which helps as a fresher.")
    else:
        feedback.append("You need at least one role-relevant project to avoid looking empty as a fresher.")
    if has_numbers:
        score += 10
        positives.append("Your resume includes numbers, which makes it more credible.")
    else:
        feedback.append("Your resume has no measurable impact. Add numbers, scale, percentages, team size, or frequency wherever truthful.")
    match_score, _, missing = estimate_match(resume_text, jd_text)
    if match_score >= 60:
        score += 15
        positives.append("Your resume language aligns reasonably well with the job description.")
    elif match_score >= 35:
        score += 8
        feedback.append("Your resume partially matches the target job, but important keywords are still missing.")
    else:
        feedback.append("Your resume is too generic for this role and will likely struggle in ATS screening.")
    if target_role.strip():
        score += 5
        positives.append("You are targeting a defined role, which makes customization easier.")
    else:
        feedback.append("You have not defined a target role clearly.")
    if profile_type == "Career switch":
        feedback.append("As a career switcher, you must highlight transferable achievements more clearly.")
    score = max(0, min(score, 100))
    if score < 55:
        feedback.insert(0, "This resume will likely get rejected in its current form for competitive roles.")
    elif score < 75:
        feedback.insert(0, "This resume is decent but still not strong enough to maximize shortlist chances.")
    else:
        feedback.insert(0, "This resume is in good shape, but a few role-specific improvements can still raise your chances.")
    if missing:
        suggestions.append(f"Add these missing job keywords where truthful: {', '.join(missing[:6])}.")
    if not has_numbers:
        suggestions.append("Rewrite at least 2 bullets with numbers or measurable outcomes.")
    if proj_lines < 1:
        suggestions.append(f"Add 1 project tailored to {target_role or 'your target role'}.")
    if exp_lines < 2:
        suggestions.append("Turn internships, college leadership, shop work, volunteering, or event coordination into experience bullets.")
    if skills_count < 5:
        suggestions.append("Expand your skills section with relevant tools and functional skills.")
    if not suggestions:
        suggestions.append("Customize this version for each company instead of using one resume everywhere.")
    return score, positives[:5], feedback[:6], suggestions[:5]


def analyze_uploaded_resume(resume_text: str, jd_text: str) -> Tuple[int, List[str], List[str], List[str], List[str]]:
    score, positives, feedback, suggestions = score_resume(
        resume_text=resume_text,
        jd_text=jd_text,
        profile_type="No experience",
        skills_text=", ".join(extract_keywords(resume_text)[:10]),
        experience_text=resume_text,
        projects_text=resume_text,
        target_role="Uploaded Resume",
    )
    _, _, missing = estimate_match(resume_text, jd_text)
    return score, positives, feedback, suggestions, missing


def suggest_roles_from_resume(resume_text: str) -> Tuple[List[str], List[str]]:
    text = clean_text(resume_text).lower()
    role_map = {
        "sales": ["Key Account Manager", "Business Development Executive", "Sales Executive"],
        "marketing": ["Marketing Executive", "Brand Associate", "Category Executive"],
        "operations": ["Operations Executive", "Supply Chain Analyst", "Process Coordinator"],
        "analysis": ["Business Analyst", "Sales Analyst", "Operations Analyst"],
        "communication": ["Client Success Executive", "Inside Sales Executive", "Account Coordinator"],
    }
    recommended = []
    for key, roles in role_map.items():
        if key in text:
            recommended.extend(roles)
    if not recommended:
        recommended = ["Business Development Executive", "Operations Executive", "Marketing Executive"]
    not_ideal = ["Highly technical software engineering roles", "Senior management roles", "Deep-specialist roles requiring years of experience"]
    return list(dict.fromkeys(recommended))[:5], not_ideal

# -------- JOB DISCOVERY --------

def slugify_for_naukri(query: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9 ]+", "", query).strip().lower()
    slug = re.sub(r"\s+", "-", slug)
    return slug

def build_job_search_queries(roles: List[str], location: str) -> List[str]:
    queries = []
    for role in roles[:5]:
        q = f"{role} fresher {location}"
        queries.append(q.strip())
    return queries

def build_job_links(queries: List[str]) -> List[dict]:
    links = []
    for q in queries:
        encoded = q.replace(" ", "%20")
        links.append({
            "query": q,
            "linkedin": f"https://www.linkedin.com/jobs/search/?keywords={encoded}",
            "indeed": f"https://www.indeed.com/jobs?q={encoded}",
            "naukri": f"https://www.naukri.com/{slugify_for_naukri(q)}-jobs",
        })
    return links

def build_apply_strategy(roles: List[str]) -> List[str]:
    return [
        "Apply to 5 startup roles (faster hiring)",
        "Apply to 5 mid-size companies",
        "Apply to 5 branded companies",
        "Focus on jobs posted in last 7 days",
    ]

def build_recruiter_message(role: str, company: str):
    return f"""Hi, I came across the {role} role at {company}.
My profile aligns well and I would love to be considered.

Would really appreciate a chance to connect. Thank you!"""
def generate_action_plan(target_role: str, missing_skills: List[str], resume_score: int, profile_type: str) -> List[str]:
    role = target_role or "your target role"
    top_missing = ", ".join(missing_skills[:3]) if missing_skills else "role-specific keywords"
    plan = [
        f"Day 1: Rewrite your resume for {role} and fix the weakest bullets first.",
        f"Day 2: Add 1 relevant project and include these missing keywords where truthful: {top_missing}.",
        "Day 3: Apply to 15 carefully selected roles instead of mass-applying blindly.",
        "Day 4: Prepare answers for 'Tell me about yourself', 'Why this role?', and one strengths question.",
        "Day 5: Improve LinkedIn headline, About section, and featured achievements.",
        "Day 6: Follow up on earlier applications and message 5 recruiters or alumni.",
        "Day 7: Review rejections, refine your resume again, and apply to another 15 better-matched roles.",
    ]
    if resume_score < 55:
        plan[0] = f"Day 1: Do not apply yet. First rebuild your resume for {role} because the current version is too weak."
    if profile_type == "No experience":
        plan[1] = f"Day 2: Add project, internship-style work, event responsibility, or family business contribution relevant to {role}."
    return plan


def parse_resume_sections(resume_text: str) -> dict:
    lines = [line.rstrip() for line in (resume_text or "").split("\n")]
    sections = {"HEADER": []}
    current = "HEADER"
    for line in lines:
        if line.strip().startswith("## "):
            current = line.replace("##", "").strip().upper()
            sections[current] = []
        else:
            sections.setdefault(current, []).append(line)
    parsed = {"name": "", "contact": "", "sections": {}}
    header_lines = [x.strip() for x in sections.get("HEADER", []) if x.strip()]
    if header_lines:
        parsed["name"] = header_lines[0].replace("#", "").strip()
    if len(header_lines) > 1:
        parsed["contact"] = header_lines[1]
    for key, value in sections.items():
        if key != "HEADER":
            parsed["sections"][key.title()] = [v.strip() for v in value if v.strip()]
    return parsed


def make_pdf_bytes(title: str, content: str, template_name: str = "Corporate Clean") -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    parsed = parse_resume_sections(content)
    y = height - 42
    left = 42
    name = parsed.get("name") or title
    contact = parsed.get("contact", "")
    pdf.setFont("Helvetica-Bold", 18)
    pdf.drawString(left, y, name)
    y -= 18
    pdf.setFont("Helvetica", 10)
    if contact:
        pdf.drawString(left, y, contact)
        y -= 20
    else:
        y -= 10
    def draw_wrapped(text_line: str, font_name: str = "Helvetica", font_size: int = 10, indent: int = 0, leading: int = 14):
        nonlocal y
        clean_line = text_line.replace("#", "").replace("*", "").strip()
        if not clean_line:
            y -= 8
            return
        words = clean_line.split()
        current = ""
        wrapped = []
        for word in words:
            test = f"{current} {word}".strip()
            if pdf.stringWidth(test, font_name, font_size) < width - (left + indent) - 42:
                current = test
            else:
                wrapped.append(current)
                current = word
        if current:
            wrapped.append(current)
        pdf.setFont(font_name, font_size)
        for line in wrapped:
            if y < 52:
                pdf.showPage()
                y = height - 42
                pdf.setFont(font_name, font_size)
            pdf.drawString(left + indent, y, line)
            y -= leading
    for section_name, section_lines in parsed.get("sections", {}).items():
        if y < 80:
            pdf.showPage()
            y = height - 42
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(left, y, section_name)
        y -= 14
        pdf.line(left, y, width - 42, y)
        y -= 10
        for line in section_lines:
            if line.startswith("- "):
                draw_wrapped(f"• {line[2:]}", indent=8)
            else:
                draw_wrapped(line)
        y -= 4
    pdf.save()
    buffer.seek(0)
    return buffer.read()


def make_docx_bytes(resume_text: str, template_name: str = "Corporate Clean") -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is not available")
    parsed = parse_resume_sections(resume_text)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)
    name = parsed.get("name") or "Candidate Name"
    contact = parsed.get("contact", "")
    if template_name == "Modern Fresher":
        header_align = WD_ALIGN_PARAGRAPH.LEFT
        name_size = 20
        section_size = 12
    elif template_name == "ATS Minimal":
        header_align = WD_ALIGN_PARAGRAPH.LEFT
        name_size = 16
        section_size = 11
    else:
        header_align = WD_ALIGN_PARAGRAPH.CENTER
        name_size = 18
        section_size = 11.5
    p = doc.add_paragraph()
    p.alignment = header_align
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(name_size)
    if contact:
        p2 = doc.add_paragraph()
        p2.alignment = header_align
        r2 = p2.add_run(contact)
        r2.font.size = Pt(10)
    for section_name, section_lines in parsed.get("sections", {}).items():
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(3)
        run = heading.add_run(section_name.upper())
        run.bold = True
        run.font.size = Pt(section_size)
        if template_name == "Corporate Clean":
            run.underline = True
        for line in section_lines:
            if line.startswith("- "):
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.paragraph_format.space_after = Pt(1)
                bullet.add_run(line[2:])
            else:
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(2)
                para.add_run(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def load_example_data() -> dict:
    return {
        "name": "Tamoghna Adhikari",
        "email": "tamoghna@example.com",
        "phone": "+91XXXXXXXXXX",
        "location": "Kolkata, India",
        "target_role": "Key Account Manager",
        "profile_type": "Internship only",
        "summary": "Recent PGDM graduate targeting sales and account management roles with strong communication, analytical thinking, and market understanding.",
        "education": "PGDM, IMT Nagpur\nB.Tech in Food Technology, Institute Name",
        "experience": "Conducted market and competitor analysis for academic and live projects.\nSupported quality and operations work during internships.\nManaged coordination and communication responsibilities in team-based assignments.",
        "projects": "Built a buyer persona and customer journey framework for a consumer brand.\nCreated a market analysis and go-to-market recommendation for a business case competition.",
        "certifications": "Advanced Excel\nGoogle Analytics Fundamentals",
        "skills": "Excel, PowerPoint, Communication, Market Research, Sales, Analysis, Presentation Skills",
        "company": "Zomato",
        "jd_text": "We are looking for a Key Account Manager to manage partner relationships, drive growth, analyze performance, coordinate with internal teams, and improve merchant success. Strong communication, problem-solving, stakeholder management, and Excel skills are preferred."
    }


def get_openai_client():
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or OpenAI is None:
        return None
    try:
        return OpenAI(api_key=api_key)
    except Exception:
        return None


def call_llm(system_prompt: str, user_prompt: str) -> str:
    client = get_openai_client()
    if client is None:
        raise RuntimeError("OpenAI client not available")
    response = client.chat.completions.create(
        model=os.getenv("OPENAI_MODEL", "gpt-4.1-mini"),
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=0.4,
    )
    return response.choices[0].message.content.strip()


def save_current_version(name: str, company: str, target_role: str):
    entry = {
        "label": f"{len(st.session_state['history']) + 1}. {target_role or 'Untitled Role'} @ {company or 'General'}",
        "name": name,
        "company": company,
        "target_role": target_role,
        "resume_output": st.session_state.get("resume_output", ""),
        "cover_output": st.session_state.get("cover_output", ""),
        "resume_score": st.session_state.get("resume_score"),
        "match_score": st.session_state.get("match_score"),
        "feedback": st.session_state.get("feedback", []),
        "suggestions": st.session_state.get("suggestions", []),
        "missing": st.session_state.get("missing", []),
        "action_plan": st.session_state.get("action_plan", []),
    }
    st.session_state["history"].append(entry)


def load_history_entry(entry: dict):
    for key in ["resume_output","cover_output","resume_score","match_score","feedback","suggestions","missing","action_plan"]:
        st.session_state[key] = entry.get(key)


# ---------- Session ----------
def ensure_state():
    defaults = {
        "history": [],
        "applications": [],
        "user_feedback_log": [],
        "example_loaded": False,
        "path_choice": "Upload Existing Resume",
        "report_unlocked": False,
        "resume_template": "Corporate Clean",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v
ensure_state()


# ---------- Sidebar ----------
with st.sidebar:
    st.markdown("## Product Controls")
    use_ai = st.toggle("Use OpenAI API", value=False)
    st.selectbox("Resume Template", ["Corporate Clean", "Modern Fresher", "ATS Minimal"], key="resume_template")
    st.markdown("---")
    st.markdown("### Saved versions")
    history = st.session_state.get("history", [])
    if history:
        selected_label = st.selectbox("Load past version", [h["label"] for h in history])
        idx = [h["label"] for h in history].index(selected_label)
        a, b = st.columns(2)
        if a.button("Load", use_container_width=True):
            load_history_entry(history[idx])
        if b.button("Delete", use_container_width=True):
            st.session_state["history"].pop(idx)
            st.rerun()
    else:
        st.caption("No saved versions yet.")
    st.markdown("---")
    st.markdown("### Value promise")
    st.markdown("- Diagnose why a fresher resume gets rejected")
    st.markdown("- Show exact fixes")
    st.markdown("- Recommend suitable roles")
    st.markdown("- Guide next steps")
    st.markdown("<div class='small'>For launch, the unlock button simulates the paid flow. Add real payment later with Razorpay or Stripe.</div>", unsafe_allow_html=True)


# ---------- Landing ----------
st.markdown(
    """
    <div class='hero'>
        <div class='hero-title'>Applied to 100 jobs and still no calls?</div>
        <div class='hero-sub'>Fix your resume, target the right roles, and get shortlisted faster. This is an AI job search coach built for freshers who need clarity, not just generic AI answers.</div>
        <div class='chip-row'>
            <div class='chip'>ATS Audit</div>
            <div class='chip'>Resume Rewrite</div>
            <div class='chip'>Role Suggestions</div>
            <div class='chip'>Interview Prep</div>
            <div class='chip'>Application Tracker</div>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    """
    <div class='step-grid'>
        <div class='step'><div class='step-num'>Step 1</div><div class='step-head'>Upload or build</div><div class='step-copy'>Start with your existing resume or build one from scratch.</div></div>
        <div class='step'><div class='step-num'>Step 2</div><div class='step-head'>See the brutal report</div><div class='step-copy'>Understand why your current resume is not converting.</div></div>
        <div class='step'><div class='step-num'>Step 3</div><div class='step-head'>Unlock exact fixes</div><div class='step-copy'>Get rewritten content, role clarity, and a concrete plan.</div></div>
        <div class='step'><div class='step-num'>Step 4</div><div class='step-head'>Apply with confidence</div><div class='step-copy'>Use interview prep and tracking to stay consistent.</div></div>
    </div>
    """,
    unsafe_allow_html=True,
)

h1, h2, h3 = st.columns([1.2, 1.2, 2.2])
with h1:
    st.radio("What do you want to do?", ["Upload Existing Resume", "Build Resume From Scratch"], key="path_choice")
with h2:
    if st.button("Load Example Data", use_container_width=True):
        for k, v in load_example_data().items():
            st.session_state[k] = v
        st.session_state["example_loaded"] = True
with h3:
    if st.session_state.get("example_loaded"):
        st.success("Example data loaded. You can test the flow immediately.")
    else:
        st.info("Start with the path selection above, then use the free audit preview below.")


# ---------- Candidate / JD input ----------
left, right = st.columns([1.02, 0.98], gap="large")

with left:
    st.markdown("<div class='soft-card'><div class='section-title'>Candidate Profile</div><div class='section-sub'>These details power resume generation and the full report after unlock.</div></div>", unsafe_allow_html=True)
    name = st.text_input("Full Name", value=st.session_state.get("name", ""), placeholder="Tamoghna Adhikari")
    email = st.text_input("Email", value=st.session_state.get("email", ""), placeholder="you@example.com")
    phone = st.text_input("Phone", value=st.session_state.get("phone", ""), placeholder="+91XXXXXXXXXX")
    location = st.text_input("Location", value=st.session_state.get("location", ""), placeholder="Kolkata, India")
    target_role = st.text_input("Target Role", value=st.session_state.get("target_role", ""), placeholder="Key Account Manager")
    profile_options = ["No experience", "Internship only", "Career switch", "Experienced"]
    profile_type = st.selectbox("Profile Type", profile_options, index=profile_options.index(st.session_state.get("profile_type", "No experience")) if st.session_state.get("profile_type", "No experience") in profile_options else 0)
    summary = st.text_area("Professional Summary", value=st.session_state.get("summary", ""), height=95)
    education = st.text_area("Education", value=st.session_state.get("education", ""), height=90)
    experience = st.text_area("Experience / Internships / Responsibilities", value=st.session_state.get("experience", ""), height=110)
    projects = st.text_area("Projects", value=st.session_state.get("projects", ""), height=90)
    certifications = st.text_area("Certifications", value=st.session_state.get("certifications", ""), height=70)
    skills = st.text_area("Skills (comma-separated)", value=st.session_state.get("skills", ""), height=80)

with right:
    st.markdown("<div class='soft-card'><div class='section-title'>Target Job</div><div class='section-sub'>Paste a real job description to improve scoring and recommendations.</div></div>", unsafe_allow_html=True)
    company = st.text_input("Company Name", value=st.session_state.get("company", ""), placeholder="Zomato")
    jd_text = st.text_area("Paste Job Description", value=st.session_state.get("jd_text", ""), height=220)

    if st.session_state["path_choice"] == "Upload Existing Resume":
        uploaded_cv = st.file_uploader("Upload your CV (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"], key="uploaded_cv")
        if uploaded_cv is not None:
            uploaded_text = extract_text_from_upload(uploaded_cv)
            st.session_state["uploaded_resume_text"] = uploaded_text
            if uploaded_text.strip():
                st.success("Resume text extracted successfully.")
            else:
                st.warning("Could not extract text properly. Try a text-based PDF, DOCX, or TXT file.")
        st.text_area("Extracted Resume Text", value=st.session_state.get("uploaded_resume_text", ""), height=150)
    else:
        st.markdown("<div class='small'>You chose to build from scratch. The free preview will use the details you entered to create a draft resume and audit it.</div>", unsafe_allow_html=True)


# ---------- Free audit preview ----------
st.markdown("---")
st.markdown("<div class='soft-card'><div class='section-title'>Free Audit Preview</div><div class='section-sub'>Show the pain first. Lock the solution until the user unlocks the full report.</div></div>", unsafe_allow_html=True)

preview_col, pay_col = st.columns([1.25, 0.75], gap="large")

with preview_col:
    analyze_clicked = st.button("Start Free Audit", use_container_width=True)
    if analyze_clicked:
        if st.session_state["path_choice"] == "Upload Existing Resume":
            resume_source = st.session_state.get("uploaded_resume_text", "")
        else:
            resume_source = build_resume_fallback(name, email, phone, location, target_role, profile_type, summary, education, experience, skills, projects, certifications)
            st.session_state["resume_output"] = resume_source

        if not clean_text(resume_source):
            st.warning("Please upload a readable CV or enter enough profile details first.")
        else:
            score, positives, feedback, suggestions, missing = analyze_uploaded_resume(resume_source, jd_text)
            rec_roles, not_ideal = suggest_roles_from_resume(resume_source)
            st.session_state["preview_source_resume"] = resume_source
            st.session_state["uploaded_resume_score"] = score
            st.session_state["uploaded_positives"] = positives
            st.session_state["uploaded_feedback"] = feedback
            st.session_state["uploaded_suggestions"] = suggestions
            st.session_state["uploaded_missing"] = missing
            st.session_state["recommended_roles"] = rec_roles
            st.session_state["not_ideal_roles"] = not_ideal

    if st.session_state.get("uploaded_resume_score") is not None:
        score = st.session_state["uploaded_resume_score"]
        if score < 55:
            st.markdown("<div class='warning-card'><b>This will likely get rejected.</b><br>Use the report below to understand the top problems.</div>", unsafe_allow_html=True)
        else:
            st.markdown("<div class='success-card'><b>This is workable but not optimized.</b><br>You can still improve your shortlist chances with the full report.</div>", unsafe_allow_html=True)
        m1, m2 = st.columns(2)
        m1.metric("Resume Score", f"{score}/100")
        m2.metric("Likely Match", f"{st.session_state.get('match_score', max(25, min(score, 90)))}%")

        st.markdown("### Top problems")
        for item in st.session_state.get("uploaded_feedback", [])[:3]:
            st.write(f"- {item}")

        st.markdown("### What is working")
        for item in st.session_state.get("uploaded_positives", [])[:2]:
            st.write(f"- {item}")

        st.markdown("### Preview of locked insights")
        st.markdown(
            """
            <div class='blur-box'>
                <div class='blur-inner'>
                    <ul>
                        <li>Exact fixes to raise the score above 75</li>
                        <li>Rewritten resume content optimized for the JD</li>
                        <li>Best roles for this profile right now</li>
                        <li>7-day application plan and interview prep</li>
                    </ul>
                </div>
                <div class='lock-overlay'>🔒 Unlock full report to see the solution</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

with pay_col:
    st.markdown(
        """
        <div class='paywall'>
            <div class='pay-title'>Unlock the full report</div>
            <div class='small'>Built for users who do not just want feedback, but want exact fixes and a plan.</div>
            <div class='pay-price'>₹299</div>
            <div class='small'>One-time unlock includes full ATS audit, rewritten resume, role suggestions, interview prep, and a 7-day action plan.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if st.button("Unlock Full Report", key="unlock_full_report", use_container_width=True):
        st.session_state["report_unlocked"] = True
        st.success("Full report unlocked for this session.")
    if st.session_state.get("report_unlocked"):
        st.info("Full report is unlocked below.")
    else:
        st.caption("For now, this simulates a paid unlock. Add payment integration later.")


# ---------- Full product experience ----------
if st.session_state.get("report_unlocked"):
    st.markdown("---")
    st.markdown("<div class='soft-card'><div class='section-title'>Full Report Dashboard</div><div class='section-sub'>This is the paid experience: fix, improve, target, prepare, and track.</div></div>", unsafe_allow_html=True)

    save_col, note_col = st.columns([1, 2])
    with save_col:
        if st.button("Save Current Version", use_container_width=True):
            save_current_version(name, company, target_role)
            st.success("Current version saved.")
    with note_col:
        st.markdown("<div class='small'>Best sequence: Resume Builder → Auto Improve → Resume Score → Job Match → 7-Day Plan → Interview → Tracker.</div>", unsafe_allow_html=True)

    tabs = st.tabs(["Resume Builder", "Auto Improve", "Resume Score", "Cover Letter", "Job Match", "7-Day Plan", "Interview", "Role Clarity", "Job Discovery"])

    with tabs[0]:
        st.markdown("### Build or review your final resume")
        if st.button("Generate Resume", key="generate_resume", use_container_width=True):
            try:
                if use_ai:
                    prompt = f"""
Create an ATS-friendly resume in markdown for this candidate.
Name: {name}
Email: {email}
Phone: {phone}
Location: {location}
Target role: {target_role}
Profile type: {profile_type}
Summary: {summary}
Education: {education}
Experience: {experience}
Projects: {projects}
Certifications: {certifications}
Skills: {skills}

Use strong action verbs. Keep it concise. Do not invent facts.
"""
                    resume_output = call_llm("You are an expert resume writer for freshers and MBA graduates.", prompt)
                else:
                    resume_output = build_resume_fallback(name, email, phone, location, target_role, profile_type, summary, education, experience, skills, projects, certifications)
                st.session_state["resume_output"] = resume_output
            except Exception as e:
                st.error(f"Could not generate resume: {e}")
        resume_output = st.session_state.get("resume_output", st.session_state.get("preview_source_resume", ""))
        st.text_area("Resume Output", value=resume_output, height=360)
        if resume_output:
            d1, d2, d3 = st.columns(3)
            d1.download_button("Download .md", data=resume_output, file_name="resume.md", mime="text/markdown", use_container_width=True)
            if REPORTLAB_AVAILABLE:
                d2.download_button("Download PDF", data=make_pdf_bytes("Resume", resume_output, st.session_state.get("resume_template", "Corporate Clean")), file_name="resume.pdf", mime="application/pdf", use_container_width=True)
            if Document is not None:
                d3.download_button("Download Word", data=make_docx_bytes(resume_output, st.session_state.get("resume_template", "Corporate Clean")), file_name="resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

    with tabs[1]:
        st.markdown("### Auto improve weak bullets")
        if st.button("Auto Improve My Resume", key="auto_improve", use_container_width=True):
            base_resume = st.session_state.get("resume_output") or st.session_state.get("preview_source_resume") or build_resume_fallback(name, email, phone, location, target_role, profile_type, summary, education, experience, skills, projects, certifications)
            try:
                if use_ai:
                    prompt = f"""
Improve this resume for a fresher targeting {target_role}.

Resume:
{base_resume}

Job description:
{jd_text}

Rewrite weak bullets, align language with the JD, and keep it truthful. Return markdown only.
"""
                    improved = call_llm("You are an expert ATS resume optimizer.", prompt)
                else:
                    improved = re.sub(r"\bhandled\b", "Managed", base_resume, flags=re.I)
                    improved = re.sub(r"\bworked on\b", "Executed", improved, flags=re.I)
                    improved = re.sub(r"\bhelped\b", "Contributed to", improved, flags=re.I)
                st.session_state["resume_output"] = improved
                st.success("Resume improved.")
            except Exception as e:
                st.error(f"Could not improve resume: {e}")
        st.text_area("Improved Resume Preview", value=st.session_state.get("resume_output", ""), height=340)

    with tabs[2]:
        st.markdown("### Resume Score Breakdown")
        if st.button("Score My Resume", key="score_resume", use_container_width=True):
            resume_base = st.session_state.get("resume_output") or st.session_state.get("preview_source_resume") or build_resume_fallback(name, email, phone, location, target_role, profile_type, summary, education, experience, skills, projects, certifications)
            score, positives, feedback, suggestions = score_resume(resume_base, jd_text, profile_type, skills, experience, projects, target_role)
            st.session_state["resume_score"] = score
            st.session_state["positives"] = positives
            st.session_state["feedback"] = feedback
            st.session_state["suggestions"] = suggestions
        score = st.session_state.get("resume_score")
        if score is not None:
            m1, m2 = st.columns(2)
            m1.metric("Resume Score", f"{score}/100")
            m2.metric("Estimated ATS Readiness", f"{max(20, min(score+5, 98))}%")
            st.markdown("#### What is working")
            for item in st.session_state.get("positives", []):
                st.write(f"- {item}")
            st.markdown("#### Brutal feedback")
            for item in st.session_state.get("feedback", []):
                st.write(f"- {item}")
            st.markdown("#### Exact fixes")
            for item in st.session_state.get("suggestions", []):
                st.write(f"- {item}")

    with tabs[3]:
        st.markdown("### Tailored cover letter")
        if st.button("Generate Cover Letter", key="cover_letter", use_container_width=True):
            try:
                if use_ai:
                    prompt = f"""
Write a concise personalized cover letter for:
Candidate: {name}
Company: {company}
Role: {target_role}
Summary: {summary}
Education: {education}
Experience: {experience}
Projects: {projects}
Skills: {skills}
Job description: {jd_text}
Do not invent facts.
"""
                    st.session_state["cover_output"] = call_llm("You are an expert cover letter writer.", prompt)
                else:
                    st.session_state["cover_output"] = build_cover_letter_fallback(name, company or "the company", target_role or "the role", summary, skills, jd_text)
            except Exception as e:
                st.error(f"Could not generate cover letter: {e}")
        cover_output = st.session_state.get("cover_output", "")
        st.text_area("Cover Letter Output", value=cover_output, height=320)
        if cover_output:
            c1, c2, c3 = st.columns(3)
            c1.download_button("Download .txt", data=cover_output, file_name="cover_letter.txt", mime="text/plain", use_container_width=True)
            if REPORTLAB_AVAILABLE:
                c2.download_button("Download PDF", data=make_pdf_bytes("Cover Letter", cover_output, st.session_state.get("resume_template", "Corporate Clean")), file_name="cover_letter.pdf", mime="application/pdf", use_container_width=True)
            if Document is not None:
                c3.download_button("Download Word", data=make_docx_bytes(cover_output, st.session_state.get("resume_template", "Corporate Clean")), file_name="cover_letter.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

    with tabs[4]:
        st.markdown("### Job match and gap diagnosis")
        if st.button("Analyze Job Match", key="job_match", use_container_width=True):
            resume_base = st.session_state.get("resume_output") or st.session_state.get("preview_source_resume", "")
            try:
                if use_ai:
                    prompt = f"Return valid JSON only with keys match_score, matched_skills, missing_skills, suggestions. Resume:\n{resume_base}\n\nJob description:\n{jd_text}"
                    raw = call_llm("You are a hiring analyst. Return only valid JSON.", prompt)
                    parsed = json.loads(raw.replace("```json", "").replace("```", "").strip())
                    st.session_state["match_score"] = int(parsed.get("match_score", 0))
                    st.session_state["matched"] = parsed.get("matched_skills", [])
                    st.session_state["missing"] = parsed.get("missing_skills", [])
                    st.session_state["match_suggestions"] = parsed.get("suggestions", [])
                else:
                    score, matched, missing = estimate_match(resume_base, jd_text)
                    st.session_state["match_score"] = score
                    st.session_state["matched"] = matched
                    st.session_state["missing"] = missing
                    st.session_state["match_suggestions"] = ["Add missing keywords from the JD where truthful.", "Make one project directly relevant to this role.", "Rewrite bullets to reflect the responsibilities more closely."]
            except Exception as e:
                st.error(f"Could not analyze match: {e}")
        match_score = st.session_state.get("match_score")
        if match_score is not None:
            st.metric("Job Match Score", f"{match_score}%")
            a, b = st.columns(2)
            with a:
                st.markdown("#### Matched skills")
                for item in st.session_state.get("matched", []):
                    st.write(f"- {item}")
            with b:
                st.markdown("#### Missing skills / keywords")
                for item in st.session_state.get("missing", []):
                    st.write(f"- {item}")
            st.markdown("#### Fix plan")
            for item in st.session_state.get("match_suggestions", []):
                st.write(f"- {item}")

    with tabs[5]:
        st.markdown("### 7-day application plan")
        if st.button("Generate 7-Day Plan", key="action_plan", use_container_width=True):
            resume_base = st.session_state.get("resume_output") or st.session_state.get("preview_source_resume", "")
            resume_score = st.session_state.get("resume_score")
            if resume_score is None:
                resume_score, _, _, _ = score_resume(resume_base, jd_text, profile_type, skills, experience, projects, target_role)
            st.session_state["action_plan"] = generate_action_plan(target_role, st.session_state.get("missing", []), resume_score, profile_type)
        for item in st.session_state.get("action_plan", []):
            st.write(f"- {item}")

    with tabs[6]:
        st.markdown("### Interview simulator")
        role_for_interview = target_role or "your target role"
        if st.button("Generate Interview Questions", key="interview_questions", use_container_width=True):
            try:
                if use_ai:
                    prompt = f"Generate 5 common interview questions for a fresher applying for {role_for_interview}. Also provide ideal structured answers."
                    st.session_state["interview_output"] = call_llm("You are an expert interview coach.", prompt)
                else:
                    st.session_state["interview_output"] = f"1. Tell me about yourself.\nAnswer: Briefly explain your education, key skills, and interest in {role_for_interview}.\n\n2. Why do you want this role?\nAnswer: Align your interest with the role and company.\n\n3. What are your strengths?\nAnswer: Mention 2 to 3 strengths with examples.\n\n4. Describe a challenge you faced.\nAnswer: Use STAR method.\n\n5. Why should we hire you?\nAnswer: Connect your skills and willingness to learn with the job."
            except Exception as e:
                st.error(f"Error generating questions: {e}")
        st.text_area("Interview Questions & Answers", value=st.session_state.get("interview_output", ""), height=240)
        answer = st.text_area("Practice Answer", key="practice_answer", height=140)
        if st.button("Evaluate My Answer", key="evaluate_interview", use_container_width=True):
            try:
                if use_ai:
                    eval_prompt = f"Evaluate this fresher interview answer for clarity, structure, confidence, and give improvement suggestions.\n\nAnswer:\n{answer}"
                    st.session_state["interview_feedback"] = call_llm("You are an interview coach.", eval_prompt)
                else:
                    st.session_state["interview_feedback"] = "Your answer is decent. Improve structure using STAR method and be more specific with examples."
            except Exception as e:
                st.error(f"Error evaluating answer: {e}")
        if st.session_state.get("interview_feedback"):
            st.markdown("#### Feedback")
            st.write(st.session_state["interview_feedback"])

        with tabs[7]:
        st.markdown("### Best roles for this profile")
        roles = st.session_state.get("recommended_roles", [])
        avoid = st.session_state.get("not_ideal_roles", [])

        if roles:
            st.markdown("#### Best-fit roles")
            for item in roles:
                st.write(f"- {item}")

        if avoid:
            st.markdown("#### Not ideal right now")
            for item in avoid:
                st.write(f"- {item}")

    with tabs[8]:
        st.markdown("### 🔎 Job Discovery")
        st.markdown("Find real job opportunities based on your target role.")

        search_role = st.text_input("Search Role", value=target_role)
        search_location = st.text_input("Location", value="India")

        if st.button("Find Jobs", use_container_width=True):
            if search_role:
                jobs = [
                    {"title": f"{search_role} at Zomato", "link": "https://www.zomato.com/careers"},
                    {"title": f"{search_role} at Swiggy", "link": "https://careers.swiggy.com"},
                    {"title": f"{search_role} at Amazon", "link": "https://www.amazon.jobs"},
                    {"title": f"{search_role} at Flipkart", "link": "https://www.flipkartcareers.com"},
                ]
                st.session_state["job_results"] = jobs
            else:
                st.warning("Please enter a role to search.")

        results = st.session_state.get("job_results", [])

        if results:
            st.markdown("#### 🎯 Suggested Jobs")
            for job in results:
                st.markdown(f"- [{job['title']}]({job['link']})")

    # Lower utilities
    st.markdown("---")
    col1, col2 = st.columns([1.1, 0.9], gap="large")
    with col1:
        st.markdown("<div class='soft-card'><div class='section-title'>📋 Job Application Tracker</div><div class='section-sub'>Track where you applied and what to follow up on next.</div></div>", unsafe_allow_html=True)
        with st.expander("Add New Application", expanded=False):
            app_company = st.text_input("Company", key="app_company")
            app_role = st.text_input("Role", key="app_role")
            app_status = st.selectbox("Status", ["Applied", "Interview", "Rejected", "Offer"], key="app_status")
            app_date = st.date_input("Date Applied", value=date.today(), key="app_date")
            follow_up = st.date_input("Follow-up Date", value=date.today(), key="follow_up")
            if st.button("Add Application", key="add_application", use_container_width=True):
                st.session_state["applications"].append({"company": app_company, "role": app_role, "status": app_status, "date": str(app_date), "follow_up": str(follow_up)})
                st.success("Application added.")
        apps = st.session_state.get("applications", [])
        if apps:
            m1, m2, m3 = st.columns(3)
            m1.metric("Total Applied", len(apps))
            m2.metric("Interviews", len([a for a in apps if a["status"] == "Interview"]))
            m3.metric("Offers", len([a for a in apps if a["status"] == "Offer"]))
            for i, app in enumerate(apps):
                with st.container(border=True):
                    a1, a2, a3, a4, a5 = st.columns([1.4, 1.2, 1, 1.1, 0.8])
                    a1.write(f"**{app['company']}**")
                    a2.write(app["role"])
                    a3.write(app["status"])
                    a4.write(app["follow_up"])
                    if a5.button("Delete", key=f"delete_app_{i}"):
                        st.session_state["applications"].pop(i)
                        st.rerun()
    with col2:
        st.markdown("<div class='soft-card'><div class='section-title'>🧠 Next Action</div><div class='section-sub'>Keep the product prescriptive so users always know what to do next.</div></div>", unsafe_allow_html=True)
        next_actions = []
        if not st.session_state.get("resume_output"):
            next_actions.append("Generate or import your final resume first.")
        if st.session_state.get("resume_score") is None:
            next_actions.append("Score your resume to see the strongest fixes.")
        if st.session_state.get("match_score") is None:
            next_actions.append("Run job match analysis against a real JD.")
        if not st.session_state.get("applications"):
            next_actions.append("Track your first 5 applications to build momentum.")
        if not next_actions:
            next_actions.append("Apply to 10 better-matched roles today using your improved resume.")
        for item in next_actions:
            st.write(f"- {item}")

    if st.session_state.get("history"):
        st.markdown("---")
        st.markdown("## Version Comparison")
        history = st.session_state["history"]
        labels = [h["label"] for h in history]
        if len(labels) >= 2:
            s1, s2 = st.columns(2)
            with s1:
                left_choice = st.selectbox("Version A", labels, key="compare_a")
            with s2:
                right_choice = st.selectbox("Version B", labels, index=min(1, len(labels)-1), key="compare_b")
            a = history[labels.index(left_choice)]
            b = history[labels.index(right_choice)]
            c1, c2 = st.columns(2)
            with c1:
                st.markdown(f"### {a['label']}")
                st.write(f"- Resume score: {a.get('resume_score', 'N/A')}")
                st.write(f"- Job match score: {a.get('match_score', 'N/A')}")
                for item in a.get("feedback", [])[:4]:
                    st.write(f"- {item}")
            with c2:
                st.markdown(f"### {b['label']}")
                st.write(f"- Resume score: {b.get('resume_score', 'N/A')}")
                st.write(f"- Job match score: {b.get('match_score', 'N/A')}")
                for item in b.get("feedback", [])[:4]:
                    st.write(f"- {item}")

    st.markdown("---")
    st.markdown("<div class='soft-card'><div class='section-title'>💬 In-App Feedback</div><div class='section-sub'>Collect launch feedback directly inside the product.</div></div>", unsafe_allow_html=True)
    feedback_name = st.text_input("Your Name (optional)", key="feedback_name")
    feedback_type = st.selectbox("Feedback Type", ["General", "Bug Report", "Feature Request", "Confusing UX", "What I liked"], key="feedback_type")
    feedback_text = st.text_area("Share your feedback", placeholder="What did you like, what felt confusing, what should be improved?", height=130, key="feedback_text")
    feedback_rating = st.slider("Overall Rating", 1, 5, 4, key="feedback_rating")
    feedback_recommend = st.radio("Would you recommend this app to a friend?", ["Yes", "No"], horizontal=True, key="feedback_recommend")
    f1, f2 = st.columns(2)
    with f1:
        if st.button("Submit Feedback", key="submit_feedback", use_container_width=True):
            entry = {"name": feedback_name.strip() or "Anonymous", "type": feedback_type, "text": feedback_text.strip(), "rating": feedback_rating, "recommend": feedback_recommend}
            if entry["text"]:
                st.session_state["user_feedback_log"].append(entry)
                st.success("Feedback submitted. Thank you.")
            else:
                st.warning("Please write some feedback before submitting.")
    with f2:
        if st.button("AI Summarize Feedback", key="summarize_feedback", use_container_width=True):
            logs = st.session_state.get("user_feedback_log", [])
            if not logs:
                st.warning("No feedback submitted yet.")
            else:
                compiled = "\n\n".join([f"Name: {x['name']}\nType: {x['type']}\nFeedback: {x['text']}" for x in logs])
                try:
                    if use_ai:
                        prompt = f"Summarize the following feedback for a job-search app. Return top recurring positives, top issues, requested features, and product priorities.\n\n{compiled}"
                        st.session_state["feedback_summary"] = call_llm("You are a product analyst.", prompt)
                    else:
                        bug_count = len([x for x in logs if x['type'] == 'Bug Report'])
                        feature_count = len([x for x in logs if x['type'] == 'Feature Request'])
                        confusing_count = len([x for x in logs if x['type'] == 'Confusing UX'])
                        like_count = len([x for x in logs if x['type'] == 'What I liked'])
                        st.session_state["feedback_summary"] = f"Feedback entries: {len(logs)}\nWhat users liked: {like_count}\nBug reports: {bug_count}\nFeature requests: {feature_count}\nConfusing UX reports: {confusing_count}\n\nPriority: Fix repeated UX confusion first, then bugs, then requested features."
                except Exception as e:
                    st.error(f"Could not summarize feedback: {e}")
    logs = st.session_state.get("user_feedback_log", [])
    if logs:
        a, b, c = st.columns(3)
        avg_rating = round(sum(x.get("rating", 0) for x in logs) / len(logs), 1)
        yes_pct = round((len([x for x in logs if x.get('recommend') == 'Yes']) / len(logs)) * 100, 1)
        a.metric("Total Feedback", len(logs))
        b.metric("Average Rating", f"{avg_rating}/5")
        c.metric("Would Recommend", f"{yes_pct}%")
        export_data = json.dumps(logs, indent=2)
        st.download_button("Export Feedback Log (.json)", data=export_data, file_name="feedback_log.json", mime="application/json")
    if st.session_state.get("feedback_summary"):
        st.markdown("### AI Feedback Summary")
        st.write(st.session_state["feedback_summary"])

else:
    st.markdown("---")
    st.markdown("<div class='small'>The product currently shows the free audit and a simulated unlock flow. Once you add a payment gateway, the full dashboard can become your paid experience.</div>", unsafe_allow_html=True)
